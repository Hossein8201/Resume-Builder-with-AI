from reportlab.lib.pagesizes import letter 
from reportlab.pdfgen import canvas
from docx import Document

def fill_template1_pdf(c, user_info, width, height):
    c.drawString(100, height - 80, "Personal Information:")
    c.drawString(120, height - 100, f"Name: {user_info[2]}")
    c.drawString(120, height - 120, f"Contact: {user_info[3]}")
    c.drawString(120, height - 140, f"Email: {user_info[4]}")
    c.drawString(120, height - 160, f"Address: {user_info[5]}")
    
    c.drawString(100, height - 200, "Education:")
    c.drawString(120, height - 220, f"University: {user_info[6]}")
    c.drawString(120, height - 240, f"Degree: {user_info[7]}")
    c.drawString(120, height - 260, f"Graduation Year: {user_info[8]}")
    
    c.drawString(100, height - 300, "Work Experience:")
    c.drawString(120, height - 320, f"Job Title: {user_info[9]}")
    c.drawString(120, height - 340, f"Company: {user_info[10]}")
    c.drawString(120, height - 360, f"Start Date: {user_info[11]}")
    c.drawString(120, height - 380, f"End Date: {user_info[12]}")
    c.drawString(120, height - 400, f"Job Description: {user_info[13]}")
    
    c.drawString(100, height - 440, "Skills:")
    c.drawString(120, height - 460, f"{user_info[14]}")
    
    c.drawString(100, height - 500, "Certifications:")
    c.drawString(120, height - 520, f"{user_info[15]}")
        
def fill_template1_word(doc, user_info):
    doc.add_heading('Resume Builder', 0) 
    
    doc.add_heading('Personal Information', level=1) 
    doc.add_paragraph(f"Name: {user_info[2]}") 
    doc.add_paragraph(f"Contact: {user_info[3]}") 
    doc.add_paragraph(f"Email: {user_info[4]}") 
    doc.add_paragraph(f"Address: {user_info[5]}") 
    
    doc.add_heading('Education', level=1) 
    doc.add_paragraph(f"University: {user_info[6]}") 
    doc.add_paragraph(f"Degree: {user_info[7]}") 
    doc.add_paragraph(f"Graduation Year: {user_info[8]}") 
    
    doc.add_heading('Work Experience', level=1) 
    doc.add_paragraph(f"Job Title: {user_info[9]}") 
    doc.add_paragraph(f"Company: {user_info[10]}") 
    doc.add_paragraph(f"Start Date: {user_info[11]}") 
    doc.add_paragraph(f"End Date: {user_info[12]}") 
    doc.add_paragraph(f"Job Description: {user_info[13]}") 
    
    doc.add_heading('Skills', level=1) 
    doc.add_paragraph(user_info[14])
    
    doc.add_heading('Certifications', level=1) 
    doc.add_paragraph(user_info[15])